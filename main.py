from docx import Document
import os

def traverse_dir(document, dir):
  for filename in os.listdir(dir):
    file_path = dir + '/' + filename
    if os.path.isdir(file_path):
      traverse_dir(document, file_path)
    else:
      if not filename.endswith('.ts'):
        continue
      if filename.endswith('.spec.ts'):
        continue
      with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()

      document.add_paragraph(content)
      document.add_paragraph('\n')

doc = Document()

base_path = ''

base_dir = os.environ.get('BASE_DIR')

if base_dir and os.path.exists(base_dir):
  traverse_dir(doc, base_dir)

doc.save('output.docx')